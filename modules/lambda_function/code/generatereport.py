import boto3
import csv
import io
import os
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import RGBColor, Inches, Pt
from datetime import datetime
import json
import logging
import ast
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Initialize logging
logger = logging.getLogger()
logger.setLevel(logging.INFO)

# Define the desired order of pillars
PILLAR_ORDER = [
    'Security',
    'Cost Optimization',
    'Reliability',
    'Operational Excellence',
    'Performance Efficiency',
    'Sustainability'
]

def lambda_handler(event, context):
    logger.info("Starting Lambda execution.")
    
    # Extract milestone_number and workload_id from the event
    milestone_number = event['milestone_number']
    workload_id = event['workload_id']
    logger.info(f"Extracted milestone_number: {milestone_number}")
    logger.info(f"Extracted workload_id: {workload_id}")
    
    # Define the S3 bucket names and file names from environment variables
    template_bucket = os.environ['TEMPLATE_BUCKET']
    template_file_name = os.environ['TEMPLATE_FILE']
    csv_bucket = os.environ['CSV_BUCKET']
    output_bucket = os.environ['DESTINATION_BUCKET']
    csv_file_name = event['csv_s3_key']
    
    logger.info(f"TEMPLATE_BUCKET: {template_bucket}")
    logger.info(f"TEMPLATE_FILE: {template_file_name}")
    logger.info(f"CSV_BUCKET: {csv_bucket}")
    logger.info(f"OUTPUT_BUCKET: {output_bucket}")
    logger.info(f"CSV_FILE_NAME: {csv_file_name}")
    
    # Build the S3 key dynamically
    s3_key = template_file_name
    logger.info(f"Resolved S3 Key: {s3_key}")
    
    # Initialize AWS clients
    s3 = boto3.client('s3')
    wa_client = boto3.client('wellarchitected')

    try:
        # Fetch milestone name
        milestone_name = get_milestone_name(wa_client, workload_id, milestone_number)
        logger.info(f"Milestone Name: {milestone_name}")
        
        # Download the CSV file from S3
        logger.info("Downloading CSV file from S3.")
        csv_object = s3.get_object(Bucket=csv_bucket, Key=csv_file_name)
        csv_content = csv_object['Body'].read().decode('utf-8')
        risks_per_pillar, high_risk_items, medium_risk_items = parse_csv_and_fetch_details(
            csv_content, wa_client, workload_id, milestone_number)
        
        # Generate the graph and get its file path
        graph_image_path = generate_risk_graph(risks_per_pillar)
        
        # Download the template file from S3
        logger.info("Downloading template document from S3.")
        local_path = '/tmp/template.docx'
        s3.download_file(template_bucket, s3_key, local_path)
        logger.info("Download successful")
        
        # Open the Word document
        document = Document(local_path)
        
        # Replace placeholders in the Word document
        replace_customer_with_milestone_name(document, milestone_name)
        replace_risk_placeholders(document, high_risk_items, medium_risk_items)
        replace_graph_placeholders(document, graph_image_path)
        
        # Save the modified document
        logger.info("Saving the modified document.")
        date_str = datetime.now().strftime("%d.%m.%Y")
        report_filename = f"{milestone_name}-{date_str}-well-architected-report.docx"
        modified_file_path = '/tmp/modified_template.docx'
        document.save(modified_file_path)
        
        # Upload the modified document to the output S3 bucket
        logger.info("Uploading the modified document to S3.")
        with open(modified_file_path, 'rb') as modified_file:
            s3.upload_fileobj(modified_file, output_bucket, report_filename)
        logger.info(f"Report successfully generated and uploaded. Filename: {report_filename}")
        
        return {
            'statusCode': 200,
            'body': json.dumps({
                'message': 'Report generated and uploaded successfully.',
                'workloadId': workload_id,
                'reportFilename': report_filename,
                'milestoneName': milestone_name,
                'milestone_number': milestone_number,
                's3Bucket': output_bucket
            })
        }
    except Exception as e:
        logger.error(f"Error during Lambda execution: {str(e)}")
        raise

def get_milestone_name(wa_client, workload_id, milestone_number):
    response = wa_client.get_milestone(WorkloadId=workload_id, MilestoneNumber=milestone_number)
    return response['Milestone']['MilestoneName']

def parse_csv_and_fetch_details(csv_content, wa_client, workload_id, milestone_number):
    logger.info("Parsing CSV and fetching details.")
    risks_per_pillar = {}
    high_risk_items = {}
    medium_risk_items = {}

    csv_reader = csv.DictReader(io.StringIO(csv_content))
    for row in csv_reader:
        question_id = row['QuestionId']
        risk = row['Risk']
        logger.info(f"Processing QuestionId: {question_id}, Risk: {risk}")

        # Parse SelectedChoices from CSV
        selected_choices = []
        if 'SelectedChoices' in row and row['SelectedChoices']:
            try:
                selected_choices = ast.literal_eval(row['SelectedChoices'])
                logger.debug(f"Selected choices for QuestionId {question_id}: {selected_choices}")
            except Exception as e:
                logger.error(f"Error parsing SelectedChoices for QuestionId {question_id}: {str(e)}")

        # Fetch answer with choices
        try:
            question_details = wa_client.get_answer(
                WorkloadId=workload_id,
                LensAlias='wellarchitected',
                QuestionId=question_id,
                MilestoneNumber=milestone_number,
            )
            logger.debug(f"Question Details for {question_id}: {json.dumps(question_details)}")
        except Exception as e:
            logger.error(f"Error fetching answer for QuestionId {question_id}: {str(e)}")
            continue

        pillar_id = format_pillar_id(question_details['Answer']['PillarId'])
        question_text = question_details['Answer']['QuestionTitle']
        logger.info(f"Question Title: {question_text}, Pillar: {pillar_id}")

        # Extract unselected choices and construct URLs
        unselected_choices = []
        if 'Choices' in question_details['Answer']:
            for choice in question_details['Answer']['Choices']:
                choice_id = choice['ChoiceId']
                if choice_id not in selected_choices:
                    # Construct the documentation URL
                    doc_url = f"https://docs.aws.amazon.com/wellarchitected/latest/framework/{choice_id}.html"
                    unselected_choices.append({'title': choice['Title'], 'url': doc_url})
        else:
            logger.warning(f"No Choices found for QuestionId {question_id}")

        # Build improvement text with hyperlinks
        if unselected_choices:
            improvement_texts = []
            for item in unselected_choices:
                improvement_texts.append(f"- {item['title']} ({item['url']})")
            improvement_text = "\n".join(improvement_texts)
        else:
            improvement_text = "No improvement plans available."

        formatted_text = (
            f"{question_text}\n"
            f"Notes: {row.get('Notes', '')}\n"
            f"Improvement Plan:\n{improvement_text}"
        )

        if risk == 'HIGH':
            high_risk_items.setdefault(pillar_id, []).append(formatted_text)
            risks_per_pillar.setdefault(pillar_id, {"HIGH": 0, "MEDIUM": 0})["HIGH"] += 1
        elif risk == 'MEDIUM':
            medium_risk_items.setdefault(pillar_id, []).append(formatted_text)
            risks_per_pillar.setdefault(pillar_id, {"HIGH": 0, "MEDIUM": 0})["MEDIUM"] += 1

    return risks_per_pillar, high_risk_items, medium_risk_items

def format_pillar_id(pillar_id):
    logger.debug(f"Formatting pillar ID: {pillar_id}")
    if pillar_id.lower() == 'costoptimization':
        return 'Cost Optimization'
    elif pillar_id.lower() == 'operationalexcellence':
        return 'Operational Excellence'
    elif pillar_id.lower() == 'performanceefficiency':
        return 'Performance Efficiency'
    else:
        return ' '.join(word.capitalize() for word in pillar_id.split())

def generate_risk_graph(risks_per_pillar):
    # Ensure pillars are in the desired order
    pillars = PILLAR_ORDER
    high_risks = [risks_per_pillar.get(pillar, {"HIGH": 0})["HIGH"] for pillar in pillars]
    medium_risks = [risks_per_pillar.get(pillar, {"MEDIUM": 0})["MEDIUM"] for pillar in pillars]

    x = range(len(pillars))
    width = 0.35

    fig, ax = plt.subplots()
    ax.bar([i - width/2 for i in x], high_risks, width, label='High Risk')
    ax.bar([i + width/2 for i in x], medium_risks, width, label='Medium Risk')
    ax.set_ylabel('Counts')
    ax.set_title('Risks by Pillar')
    ax.set_xticks(list(x))
    ax.set_xticklabels(pillars, rotation=45)
    ax.legend()
    fig.tight_layout()

    graph_path = '/tmp/risk_graph.png'
    plt.savefig(graph_path)
    plt.close(fig)
    return graph_path

def replace_graph_placeholders(document, graph_path):
    for paragraph in document.paragraphs:
        if '{{pillargraph}}' in paragraph.text:
            paragraph.clear()
            paragraph.add_run().add_picture(graph_path, width=Inches(6))

def replace_risk_placeholders(document, high_risk_items, medium_risk_items):
    # Prepare high risk items text grouped by pillar
    high_risk_text = ""
    for pillar in PILLAR_ORDER:
        if pillar in high_risk_items:
            high_risk_text += f"\n{pillar}\n"
            for item in high_risk_items[pillar]:
                high_risk_text += f"{item}\n\n"  # Added extra newline for readability

    # Prepare medium risk items text grouped by pillar
    medium_risk_text = ""
    for pillar in PILLAR_ORDER:
        if pillar in medium_risk_items:
            medium_risk_text += f"\n{pillar}\n"
            for item in medium_risk_items[pillar]:
                medium_risk_text += f"{item}\n\n"  # Added extra newline for readability

    # Replace placeholders in the document
    for paragraph in document.paragraphs:
        if '{{highrisk}}' in paragraph.text:
            replace_paragraph_with_text(paragraph, high_risk_text.strip())
        elif '{{mediumrisk}}' in paragraph.text:
            replace_paragraph_with_text(paragraph, medium_risk_text.strip())

def replace_paragraph_with_text(paragraph, text):
    paragraph.clear()
    add_hyperlinked_text(paragraph, text)

def add_hyperlinked_text(paragraph, text):
    # Split the text by lines
    lines = text.split('\n')
    for line in lines:
        if '(' in line and ')' in line:
            # Extract the URL
            start = line.find('(')
            end = line.find(')', start)
            if start != -1 and end != -1:
                url = line[start+1:end]
                display_text = line[:start].strip(' -')
                # Add hyperlink
                add_hyperlink(paragraph, url, display_text)
                paragraph.add_run('\n')
        else:
            paragraph.add_run(line + '\n')

def add_hyperlink(paragraph, url, text):
    # This function adds a hyperlink to a paragraph.
    # Reference: https://stackoverflow.com/a/42334914
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Style for the hyperlink
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    
    # Text for the hyperlink
    text_element = OxmlElement('w:t')
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def replace_customer_with_milestone_name(document, milestone_name):
    for paragraph in document.paragraphs:
        if 'CUSTOMER' in paragraph.text:
            paragraph.text = paragraph.text.replace('CUSTOMER', milestone_name)
            for run in paragraph.runs:
                run.font.size = Pt(20)
